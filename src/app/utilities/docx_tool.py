""" Module that controls the involvement of creating and filling in a docx (MS Word Document)"""


from pathlib import Path
from typing import Any
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt


@dataclass
class DocxConfig:
    input_filename_path: str
    output_docx_filename: str
    title: str = "OCR Output"
    output_font_size: int = 12 
    font_style: str = field(default="Times New Roman")


class DocxTool:
    """ 
    Tool for creating and fillout out .docx documents
    
    Currently very minimal, looking to advance afterwards.

    Following a docx tutorial on youtube and jugging through documentation is not bad!
    
    """

    def __init__(self, cfg: DocxConfig = DocxConfig()) -> None:
        self._cfg = cfg


    def render_document(self, ocr_tagged: dict[str, Any], out_path: Path) -> Path:
        """ Private method to create a document"""
        doc = Document()

        doc.add_heading(self.cfg.title, level=1)

        pages: list[dict[str, Any]] = ocr_tagged.get("pages", [])
        for page in pages:
            doc.add_heading(f"Page {page.get('page_index', '?')}", level=2)

            for blk in page.get("blocks", []):
                text = (blk.get("text") or "").strip()
                if not text:
                    continue

                p = doc.add_paragraph()
                run = p.add_run(text)

                run.font.name = self._cfg.font_name
                run.font.size = Pt(self._cfg.font_size_pt)

                if blk.get("is_math"):
                    run.font.name = self._cfg.math_font_name
                    run.font.size = Pt(self._cfg.math_font_size_pt)

            doc.add_page_break()

        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        return out_path

